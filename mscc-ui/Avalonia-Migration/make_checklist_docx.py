"""Generate printable MSCC Avalonia control/tab checklist Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13 if level == 2 else 11,
            bold=True,
        )
    return p


def add_para(text, bold=False, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_check(text, status="TODO", note=""):
    mark = {"DONE": "[X]", "PARTIAL": "[~]", "TODO": "[ ]", "N/A": "[-]"}.get(
        status, "[ ]"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"{mark}  {text}")
    set_run_font(run, size=10, bold=(status == "DONE"))
    if note:
        run2 = p.add_run(f"  — {note}")
        set_run_font(run2, size=9, color=(80, 80, 80))
    return p


def add_blank_notes(n=3):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("_" * 92)
        set_run_font(run, size=9, color=(160, 160, 160))


# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MSCC Avalonia — Control & Tab Checklist")
set_run_font(r, size=18, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(
    "Printable work list for Linux / RPi GUI  ·  Client 0.4.1 baseline  ·  MSCC.Core shared with WPF"
)
set_run_font(r, size=9, color=(70, 70, 70))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Date: ______________    Tester: ______________    Display: 1024×600 or ______________"
)
set_run_font(r, size=9)

add_heading_custom("How to use this list", 1)
add_para(
    "Mark each item as you decide and implement it. Status boxes: "
    "[X] = Done (live, good enough for now) · [~] = Partial (UI only / needs polish) · "
    "[ ] = Todo · [-] = Not applicable / deferred."
)
add_para(
    "Write layout notes in the blank lines under each section (move, resize, hide, rename)."
)
add_para(
    "Suggested later order: (1) layout notes while operating, (2) wire daily operate controls, "
    "(3) spectrum polish, (4) S-meter / FWD / ALC faces (hard), (5) cal tabs last."
)

add_heading_custom("1. High-level roadmap (from here)", 1)
phases = [
    ("A. Layout & UX study (you)", "Use 0.4.1 daily; annotate this printout for moves/resizes/fonts/tabs."),
    ("B. Wire operate essentials", "Filters for real, volumes P/D, RIT, PTT/TUN carefully, ATT/AMP, MON."),
    ("C. Spectrum polish", "Zoom, S/W popup, palettes, dB cal, performance on Pi."),
    ("D. Meters (hard)", "Analog S-meter + FWD/SWR face; ALC face; HOLD/PEAK; TX vs RX swap."),
    ("E. RX helpers", "NB / NR / AN + sliders; AGC; compression."),
    ("F. CW tab", "Speed, pitch, keyer, QSK, phones — full CW page."),
    ("G. Favorites", "Client-side memories (no server required)."),
    ("H. Cal tabs", "QRP / AMP / RX IQ / TX IQ / Freq cal — after operate is solid."),
    ("I. Settings", "Appearance, COM/audio on Pi story, local port, paths."),
    ("J. Packaging", ".deb install, menu entry, version stamp."),
]
for letter, desc in phases:
    add_check(f"{letter}: {desc}", "TODO")
add_para("My notes on priority order:", bold=True)
add_blank_notes(4)

add_heading_custom("2. Already working (baseline — leave checked unless regresses)", 1)
add_para("Re-verify after big layout changes.", size=9)
baseline = [
    ("Connect / Disconnect (connect-only)", "DONE", "Does not launch servers; use MSCC Start on Pi"),
    ("Host + Port fields", "DONE", ""),
    ("UDP keep-alive + packet stats", "DONE", ""),
    ("Spectrum live pan (0xD5)", "DONE", "Frames increasing"),
    ("Waterfall scroll", "DONE", "MVP palette/dB cal"),
    ("Click-to-tune on spectrum", "DONE", ""),
    ("VFO A display (MHz)", "DONE", ""),
    ("VFO A mouse-wheel tune", "DONE", "Step from left / double-click VFO"),
    ("Step cycle (left FILTERS/STEP)", "DONE", ""),
    ("Band buttons 160…10", "DONE", "Default edge freqs"),
    ("Mode USB / LSB / CW / AM / DIG-U", "DONE", "Right rail MODE section"),
    ("Core / FW version reports", "DONE", "When server sends"),
    ("S-meter numeric text", "PARTIAL", "Digital only — not analog face"),
    ("ALC numeric text", "PARTIAL", "Digital if reported — not analog face"),
    ("Activity log + LOG clear", "DONE", ""),
    ("Three-column shell + scroll rails", "DONE", "1024×600 usable"),
    ("Tab strip (compact)", "DONE", "Headers short; content mostly empty"),
]
for t, s, n in baseline:
    add_check(t, s, n)

add_heading_custom("3. Left rail — SERVER", 1)
for t, s, n in [
    ("Connect / Disconnect button", "DONE", ""),
    ("Status line text", "DONE", ""),
    ("Auto Start checkbox", "TODO", "Currently disabled placeholder"),
    ("Launch Servers checkbox", "TODO", "Pi: usually stay OFF / document only"),
    ("Server host TextBox", "DONE", ""),
    ("Server port NumericUpDown", "DONE", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (SERVER):", bold=True)
add_blank_notes(2)

add_heading_custom("4. Left rail — Phones (operator audio)", 1)
for t, s, n in [
    ("Phones Volume slider", "TODO", "Core: speaker volume"),
    ("Phones Volume readout", "TODO", ""),
    ("Phones Mic Gain slider", "TODO", "Core: mic volume"),
    ("Phones Mic Gain readout", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (Phones):", bold=True)
add_blank_notes(2)

add_heading_custom("5. Left rail — Digital audio", 1)
for t, s, n in [
    ("Digital Volume slider", "TODO", ""),
    ("Digital Volume readout", "TODO", ""),
    ("Digital Mic Gain slider", "TODO", ""),
    ("Digital Mic Gain readout", "TODO", ""),
    ("Audio path button Phones ↔ Digital", "TODO", "Toggle P/D path"),
]:
    add_check(t, s, n)
add_para("Layout notes (Digital):", bold=True)
add_blank_notes(2)

add_heading_custom("6. Left rail — FILTERS / STEP", 1)
for t, s, n in [
    ("Lo cut button (cycle)", "PARTIAL", "UI label only — need SetFilterLowAsync"),
    ("Hi cut button (cycle)", "PARTIAL", "UI label only — need SetFilterHighAsync"),
    ("CW filter button (cycle)", "PARTIAL", "UI label only — need SetCwFilterAsync"),
    ("Step button (cycle)", "DONE", "Live for VFO wheel / spectrum"),
]:
    add_check(t, s, n)
add_para("Layout notes (Filters/Step):", bold=True)
add_blank_notes(2)

add_heading_custom("7. Left rail — Spectrum zoom & RIT", 1)
for t, s, n in [
    ("Spectrum Zoom slider 1–4×", "TODO", "Display zoom only (client viewport)"),
    ("Spectrum Zoom readout", "TODO", ""),
    ("RIT checkbox", "TODO", ""),
    ("RIT Offset slider", "TODO", ""),
    ("RIT Offset readout", "TODO", ""),
    ("Clear RIT button", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (Zoom/RIT):", bold=True)
add_blank_notes(2)

add_heading_custom("8. Left rail — Temps / status", 1)
for t, s, n in [
    ("Proficio temperature display", "PARTIAL", "Wires if server reports"),
    ("PA temperature display", "PARTIAL", ""),
    ("PA current display", "PARTIAL", ""),
    ("Packet / keep-alive / spectrum counters", "DONE", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (Temps):", bold=True)
add_blank_notes(2)

add_heading_custom("9. Center top — Meters & VFO (HARD: meters)", 1)
add_para(
    "S-meter / FWD power and ALC are the last visually hard pieces "
    "(custom drawing, needle physics, TX face swap, HOLD/PEAK). Digital placeholders exist today."
)

add_heading_custom("9a. S-meter / FWD / SWR face", 2)
for t, s, n in [
    ("S-meter digital readout (dBm text)", "PARTIAL", "Live-ish from server"),
    ("Analog S-meter face (needle + scale)", "TODO", "HARD — match WPF AnalogSMeterControl"),
    ("HOLD checkbox / behavior", "TODO", ""),
    ("PEAK needle / peak hold", "TODO", ""),
    ("TX face: FWD power 0–10 scale", "TODO", "HARD — external SWR path on WPF"),
    ("TX face: SWR digital / fault", "TODO", "Optional external SWR service"),
    ("Swap S-meter ↔ FWD face on TX", "TODO", "Depends on TX state"),
]:
    add_check(t, s, n)
add_para("Notes / sketch ideas (S-meter):", bold=True)
add_blank_notes(4)

add_heading_custom("9b. ALC meter", 2)
for t, s, n in [
    ("ALC digital readout", "PARTIAL", "If server reports ALC"),
    ("Analog ALC face (needle + scale)", "TODO", "HARD — match WPF AnalogAlcMeterControl"),
    ("ALC HOLD / PEAK options", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Notes / sketch ideas (ALC):", bold=True)
add_blank_notes(3)

add_heading_custom("9c. VFO A / VFO B", 2)
for t, s, n in [
    ("VFO A frequency display", "DONE", ""),
    ("VFO A mode display", "DONE", ""),
    ("VFO A step display", "DONE", ""),
    ("VFO A wheel tune", "DONE", ""),
    ("VFO A double-click = cycle step", "DONE", ""),
    ("VFO A click to select (split prep)", "TODO", "When VFO B live"),
    ("VFO B frequency display", "TODO", "Placeholder only now"),
    ("VFO B mode display", "TODO", ""),
    ("VFO B select / A=B / swap", "TODO", "CMD_SET_VFO etc."),
    ("Direct MHz entry on VFO (optional)", "TODO", "Removed center bar — add if needed"),
]:
    add_check(t, s, n)
add_para("Layout notes (VFO):", bold=True)
add_blank_notes(2)

add_heading_custom("10. Center — Band bar", 1)
for t, s, n in [
    ("Radio model button (Proficio / Geminus)", "TODO", "Gray-out LF/HF bands"),
    ("Band 160", "DONE", ""),
    ("Band 80", "DONE", ""),
    ("Band 60", "DONE", ""),
    ("Band 40", "DONE", ""),
    ("Band 30", "DONE", ""),
    ("Band 20", "DONE", ""),
    ("Band 17", "DONE", ""),
    ("Band 15", "DONE", ""),
    ("Band 12", "DONE", ""),
    ("Band 10", "DONE", ""),
    ("Band 2200 (Geminus)", "TODO", "Not on bar yet"),
    ("Band 630 (Geminus)", "TODO", "Not on bar yet"),
    ("GEN / USER rotate", "TODO", "WWV/CHU or Geminus cal carriers"),
    ("Active band highlight styling", "TODO", "Visual only"),
    ("Last-used freq/mode per band", "TODO", "Client INI"),
]:
    add_check(t, s, n)
add_para("Layout notes (Band bar):", bold=True)
add_blank_notes(3)

add_heading_custom("11. MAIN tab — Spectrum / waterfall", 1)
for t, s, n in [
    ("Spectrum line + fill", "DONE", "MVP renderer"),
    ("Waterfall history", "DONE", ""),
    ("Center tune line", "DONE", ""),
    ("Filter passband sketch", "PARTIAL", "Simple defaults by mode"),
    ("Click-to-tune", "DONE", ""),
    ("dB grid / labels", "PARTIAL", "Simple grid; full WPF S/W later"),
    ("dB cal (SPECTRUM_DB_OFFSET)", "PARTIAL", "Hardcoded −91.3 like WPF center"),
    ("Spectrum zoom 1–4×", "TODO", "Hook left slider"),
    ("S/W settings popup (colors, palettes, WF gain/zero)", "TODO", "Right-rail S/W button"),
    ("Peak marker", "TODO", ""),
    ("Waterfall time markers", "TODO", ""),
    ("Pan resolution 800/1600/3200", "TODO", "Match server"),
    ("Pi performance (FPS / throttle)", "PARTIAL", "Every 2nd frame now"),
    ("Freq scale labels under display", "PARTIAL", "Ticks only"),
]:
    add_check(t, s, n)
add_para("Layout notes (Spectrum size/placement):", bold=True)
add_blank_notes(4)

add_heading_custom("12. Right rail — OPERATE", 1)
for t, s, n in [
    ("PTT button", "TODO", "TX critical — careful testing"),
    ("TUN button", "TODO", "Tune carrier"),
    ("ATT button", "TODO", ""),
    ("AMP button (QRP/QRO)", "TODO", "PA bypass / full"),
    ("CMP button (compression on)", "TODO", ""),
    ("S/W button (spectrum settings)", "TODO", "Opens settings popup"),
]:
    add_check(t, s, n)
add_para("Layout notes (Operate):", bold=True)
add_blank_notes(2)

add_heading_custom("13. Right rail — MODE", 1)
for t, s, n in [
    ("USB", "DONE", ""),
    ("LSB", "DONE", ""),
    ("CW", "DONE", ""),
    ("AM", "DONE", ""),
    ("DIG-U", "DONE", "Profile; RF = USB"),
    ("FM", "TODO", "Placeholder stub — may never ship"),
    ("Mode button selected/latched styling", "TODO", "Highlight active mode"),
    ("TUNE mode (from TUN path)", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (Mode):", bold=True)
add_blank_notes(2)

add_heading_custom("14. Right rail — Compression / AGC / RX helpers", 1)
for t, s, n in [
    ("Compression slider", "TODO", ""),
    ("Compression readout", "TODO", ""),
    ("AGC Fast Release slider", "TODO", ""),
    ("AGC Fast Release readout", "TODO", ""),
    ("NB enable button", "TODO", ""),
    ("NR enable button", "TODO", ""),
    ("AN (auto notch) button", "TODO", ""),
    ("NB Pulse slider", "TODO", ""),
    ("NB Threshold slider", "TODO", ""),
    ("NR Level slider", "TODO", ""),
    ("MON (monitor) button", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (RX helpers):", bold=True)
add_blank_notes(2)

add_heading_custom("15. Right rail — CW strip (also full CW tab later)", 1)
for t, s, n in [
    ("CW Speed − button", "TODO", ""),
    ("CW Speed display", "TODO", ""),
    ("CW Speed + button", "TODO", ""),
    ("CW Pitch button / cycle", "TODO", ""),
]:
    add_check(t, s, n)
add_para("Layout notes (CW strip):", bold=True)
add_blank_notes(2)

add_heading_custom("16. Right rail — LOG / versions", 1)
for t, s, n in [
    ("LOG button (clear activity)", "DONE", "Could also open full log window"),
    ("Client version text", "DONE", ""),
    ("Core version text", "PARTIAL", "When reported"),
    ("FW version text", "PARTIAL", "When reported"),
    ("Activity list on right rail", "DONE", "Optional: move later"),
]:
    add_check(t, s, n)
add_para("Layout notes (LOG/versions):", bold=True)
add_blank_notes(2)

add_heading_custom("17. Tabs (each tab as its own work item)", 1)
add_para(
    "Header label today → full WPF name. Content status separate from “tab exists.”",
    size=9,
)
tabs = [
    ("MAIN", "Main operate + spectrum", "PARTIAL", "Spectrum + shell live; polish ongoing"),
    ("CW", "CW keyer page", "TODO", "Empty placeholder"),
    ("RX/TX", "RX/TX power, NB/NR/AN, AGC", "TODO", "Empty; some controls on right rail"),
    ("FAV", "Favorites", "TODO", "Empty"),
    ("QRP", "QRP / TRANS CAL", "TODO", "Empty; AMP gate rules later"),
    ("AMP", "AMP CAL", "TODO", "Empty; AMP-on gate later"),
    ("RX IQ", "RX I/Q balance", "TODO", "Empty"),
    ("TX IQ", "TX I/Q balance", "TODO", "Empty; QRP-only rules"),
    ("F CAL", "Frequency calibration", "TODO", "Empty"),
    ("SET", "Settings", "PARTIAL", "Local RX port only; expand later"),
]
for short, full, s, n in tabs:
    add_check(f"Tab [{short}] — {full}", s, n)

add_heading_custom("17a. CW tab content (when ready)", 2)
for t in [
    "Keyer mode",
    "Spacing",
    "Paddle",
    "Weight",
    "WPM",
    "TX hold",
    "QSK",
    "Phones",
    "Pitch list",
    "Any other CW-only controls from WPF",
]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17b. RX/TX tab content (when ready)", 2)
for t in [
    "Main / Tune / CW / SSB / AM power sliders",
    "Full power / QRP",
    "Default filters block",
    "AGC level",
    "NB/NR/AN (if moved from right rail)",
    "Any other RX/TX controls from WPF",
]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17c. Favorites tab", 2)
for t in ["List of favorites", "Save / Recall / Delete", "Per-band storage", "Name editor"]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17d. QRP CAL tab", 2)
for t in ["Band select", "Power slider", "Tune on/off", "Status lamps per band", "Commit / defaults"]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17e. AMP CAL tab", 2)
for t in ["Band select", "Amp power path", "Manual pot cal", "Status lamps", "AMP-on gate"]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17f. RX IQ / TX IQ tabs", 2)
for t in ["Band select", "Offset slider", "TX on for TX IQ", "Apply / commit", "Reset all bands"]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17g. FREQ CAL tab", 2)
for t in [
    "LOOSE / AUTO / MANUAL / CHECK / RESET",
    "PPM ±",
    "Progress / status",
    "Mini spectrum if needed",
]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("17h. SETTINGS tab", 2)
for t in [
    "UI appearance / chrome colors",
    "COM / CAT panel (Pi may use init-gui)",
    "Audio device panel (Pi: often mscc-init-gui)",
    "SWR meter settings",
    "Client INI paths (Linux vs Windows)",
    "Local RX port 8889",
    "Config reset",
]:
    add_check(t, "TODO")
add_blank_notes(2)

add_heading_custom("18. Global / cross-cutting", 1)
for t, s, n in [
    ("Linux config paths (~/.config or ~/.local)", "PARTIAL", "Logs under MSCC-Avalonia"),
    ("MSCC_Client.ini parity with WPF", "TODO", ""),
    ("Single-instance mutex", "TODO", ""),
    ("Window size/position restore", "TODO", ""),
    ("Gold button selected-state styling", "TODO", ""),
    ("Disable TX controls until connected", "PARTIAL", "Many stubs always disabled"),
    ("Session reject second client message", "TODO", "Show clearly in UI"),
    ("Keep-alive lost dialog / Continue", "PARTIAL", "Status line only now"),
    (".deb package + desktop menu", "TODO", ""),
    ("Self-contained vs framework-dependent publish", "PARTIAL", "Both exist; document one"),
    ("Windows Avalonia same binary later", "TODO", "Long-term unify"),
]:
    add_check(t, s, n)
add_para("Global notes:", bold=True)
add_blank_notes(3)

add_heading_custom("19. Your GUI change list (fill in while studying)", 1)
add_para(
    "Use this section for moves, resizes, fonts, colors, hide/show — not necessarily protocol work."
)
add_para(
    'Examples: “tab font +1”, “right rail narrower”, “S-meter taller”, “hide activity log”.',
    size=9,
)

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, h in enumerate(["#", "Change (what / where)", "Priority H/M/L", "Done?"]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for run in p.runs:
            set_run_font(run, size=9, bold=True)

for i in range(1, 21):
    row = table.add_row().cells
    row[0].text = str(i)
    row[1].text = ""
    row[2].text = ""
    row[3].text = "[ ]"
    for c in row:
        for p in c.paragraphs:
            for run in p.runs:
                set_run_font(run, size=9)

for row in table.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(5.0)
    row.cells[2].width = Inches(1.0)
    row.cells[3].width = Inches(0.7)

doc.add_paragraph()
add_para("Extra notes / sketch space:", bold=True)
add_blank_notes(8)

add_heading_custom("20. Recommended next coding sessions (when you are ready)", 1)
for t, s, n in [
    ("Session: Filters real (Lo/Hi/CW → Core opcodes)", "TODO", "High value, low risk"),
    ("Session: P/D volume + mic gain", "TODO", "High daily-use value"),
    ("Session: Mode latched styling + band highlight", "TODO", "Polish"),
    ("Session: Spectrum zoom + S/W basics", "TODO", ""),
    ("Session: Digital S-meter polish then analog face", "TODO", "HARD — meters"),
    ("Session: ALC analog face", "TODO", "HARD — meters"),
    ("Session: PTT/TUN with safety checks", "TODO", "Careful on air"),
]:
    add_check(t, s, n)

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run(
    "— End of checklist · Avalonia Migration / MSCC · Keep with printouts of WPF main tab for reference —"
)
set_run_font(r, size=8, color=(100, 100, 100))

out = Path(
    r"C:\Users\n8vet\OneDrive\Documents\MSCC-Grok-Build\Avalonia Migration\MSCC-Avalonia-Control-Checklist.docx"
)
doc.save(out)
print("Wrote", out)
print("Size", out.stat().st_size)

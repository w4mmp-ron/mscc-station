"""Generate INSTALL-FOR-PI.docx from current install narrative (matches INSTALL-FOR-PI.md)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

out = r"C:\Users\Ron\.grok\worktrees\mscc-deb\INSTALL-FOR-PI.docx"

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

for i in range(1, 4):
    hs = doc.styles[f"Heading {i}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(18)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    shd.set(qn("w:val"), "clear")
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd)
    return p


def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D6E3F0")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t


def bold_run(p, text):
    r = p.add_run(text)
    r.bold = True
    return r


# --- Title ---
doc.add_heading("MSCC on Raspberry Pi", level=1)

sub = doc.add_paragraph()
run = sub.add_run("Install in plain English")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
bold_run(p, "For: ")
p.add_run("Raspberry Pi 4 or 5, 64-bit Raspberry Pi OS (desktop recommended).")

p = doc.add_paragraph()
bold_run(p, "Goal: ")
p.add_run(
    "Install MSCC, configure it, and start/stop the servers — mostly point-and-click after package install."
)

doc.add_paragraph(
    "Current package examples (use the real filenames you were given):"
)
add_table(
    ["Package", "Example filename", "Order"],
    [
        ["PortAudio", "mscc-portaudio_19.8.2_arm64.deb", "1st"],
        ["Main stack", "mscc_1.0.27_arm64.deb", "2nd"],
        ["Setup wizard", "mscc-init-gui_1.0.10_all.deb", "3rd (recommended)"],
    ],
    [1.6, 3.8, 1.6],
)

# --- Big picture ---
doc.add_heading("Big picture (read this once)", level=2)

add_table(
    ["Where", "What runs"],
    [
        [
            "Raspberry Pi",
            "Three servers (radio + audio stack) and setup tools",
        ],
        [
            "Your PC",
            "MSCC client (screen, knobs, spectrum) — how you operate",
        ],
    ],
    [2.0, 5.0],
)

doc.add_paragraph(
    "The Pi does not need the MSCC client window. That is normal: headless servers on the Pi, GUI on the PC."
)
add_code(
    "PC (MSCC client)  ← network →  Pi (ms-sdr + sdrcore-recv + sdrcore-trans)  ← USB →  Proficio"
)

doc.add_paragraph("Two audio worlds on the Pi:")
add_table(
    ["Path", "What you hear / use", "Devices"],
    [
        [
            "Operator (phones)",
            "Headphones / mic at the Pi",
            "Real sound card — see Operator audio hardware below. Any sample rate is OK (servers adapt).",
        ],
        [
            "Digital (WSJT etc.)",
            "Digi app audio on the Pi",
            "VirtualA / VirtualB (software cables — not a physical cable)",
        ],
    ],
    [1.8, 2.2, 3.0],
)

doc.add_heading("Operator audio hardware (important on Pi 4)", level=3)
doc.add_paragraph(
    "The Proficio radio always uses USB audio for I/Q. Operator phones and mic should preferably use a "
    "different kind of path so the Pi is not running two busy USB sound devices at once."
)
add_table(
    ["Operator phones / mic", "Pi 4", "Pi 5", "Notes"],
    [
        [
            "I²S audio HAT (e.g. AudioInjector, Codec Zero, HiFiBerry with ADC+DAC)",
            "Recommended",
            "Recommended",
            "Proven, stable approach",
        ],
        [
            "Pi headphone jack / onboard audio",
            "OK for listen",
            "OK",
            "Limited mic options",
        ],
        [
            "Digital only (VirtualA / VirtualB)",
            "Recommended for digi",
            "Recommended",
            "No second physical sound card",
        ],
        [
            "Second USB headset or USB sound dongle (with Proficio also on USB)",
            "Not recommended",
            "Often OK",
            "On Pi 4, under load this can interrupt USB audio; spectrum and phones may freeze while servers still look “running”",
        ],
    ],
    [2.4, 1.2, 1.0, 2.4],
)
doc.add_paragraph(
    "In short: On Raspberry Pi 4, choose an I²S HAT or the Pi jack for operator audio, or work digi via "
    "VirtualA/B. Reserve a USB headset plus Proficio mainly for Pi 5, or accept that dual USB audio is less "
    "reliable on Pi 4."
)
doc.add_paragraph(
    "This is a supported configuration choice, not a defective install: Pi 4 is fully supported when operator "
    "audio is not a second USB sound device."
)

# --- What you need ---
doc.add_heading("What you need", level=2)
for t in [
    "Pi 4/5 with Raspberry Pi OS desktop (menu + Terminal).",
    "The three .deb files on the Pi (USB stick, Downloads folder, etc.).",
    "Optional helper: install-mscc.sh (same folder as the main .deb).",
    "Network on the Pi the first time (so apt can pull libraries).",
    "Multus / Proficio powered on when you configure (recommended).",
    "PC and Pi on the same network (or as your site requires) for the MSCC client.",
    "Operator audio plan (see table above) — especially on Pi 4, plan an I²S HAT or digi path rather than a USB headset next to the Proficio.",
]:
    doc.add_paragraph(t, style="List Number")

# --- Terminal ---
doc.add_heading("Important: use Terminal from the Desktop", level=2)
doc.add_paragraph(
    "Package install still needs a short command line. Do this on the Pi:"
)
for t in [
    "Log into the Raspberry Pi desktop (monitor or usual desktop session).",
    "Open Terminal from the desktop menu.",
    "Run the install commands in that Terminal.",
]:
    doc.add_paragraph(t, style="List Number")

doc.add_paragraph(
    "You only need enough CLI to type a few lines and enter your password. "
    "After install, use the menu:"
)
p = doc.add_paragraph()
bold_run(p, "Sound & Video")
p.add_run(" → MSCC Init, MSCC Start, MSCC Stop")
doc.add_paragraph(
    "(If the menu layout changes, search the applications menu for those three names.)"
)

# --- Install ---
doc.add_heading("Install steps", level=2)

doc.add_heading("1) Install PortAudio for MSCC (required first)", level=3)
doc.add_paragraph(
    "In Desktop Terminal, go to the folder with the .deb files:"
)
add_code("cd ~/Downloads\n# or:  cd /media/pi/YOUR_USB_NAME")

doc.add_paragraph(
    "MSCC digi (VirtualA / VirtualB) needs a special PortAudio build (Pulse + ALSA). "
    "The stock Raspberry Pi OS library alone is often not enough."
)
add_code("sudo apt install -y ./mscc-portaudio_19.8.2_arm64.deb")
doc.add_paragraph("(Use the real filename if the version number differs.)")
doc.add_paragraph(
    "This installs under /usr/local/lib. MSCC servers are built to use that copy."
)

p = doc.add_paragraph()
bold_run(p, "Quick check (recommended):")
add_code(
    "ldconfig -p | grep portaudio\n"
    "# good: a line with /usr/local/lib/libportaudio.so.2"
)
doc.add_paragraph("After the main package is installed, also check:")
add_code(
    "ldd $HOME/mscc/sdrcore-recv | grep portaudio\n"
    "# good: /usr/local/lib/...\n"
    "# bad:  only /lib/aarch64-linux-gnu/...  → digi will often fail"
)
doc.add_paragraph(
    "Do not rely on old $HOME/portaudio-install + .bashrc exports for a normal install. Use this package."
)

doc.add_heading("2) Install the main MSCC package", level=3)
p = doc.add_paragraph()
bold_run(p, "Helper script (if you have it):")
add_code("chmod +x install-mscc.sh\n./install-mscc.sh")

p = doc.add_paragraph()
bold_run(p, "Or install the .deb directly:")
add_code("sudo apt update\nsudo apt install -y ./mscc_1.0.27_arm64.deb")
doc.add_paragraph("(Use the real filename if the version differs.)")

for t in [
    "Enter your password when asked.",
    "If it asks to continue / stop old servers → type y when ready.",
    "Install does not run the setup wizard for you — that is the next step.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph(
    'You may see a technical “Notice” about _apt. '
    "If the install finished and printed “MSCC installed” / “done”, ignore that notice."
)

doc.add_paragraph("This package installs:")
add_table(
    ["Piece", "What it is"],
    [
        ["Servers", "ms-sdr, sdrcore-recv, sdrcore-trans → $HOME/mscc/"],
        ["Start/stop", "mscc command and desktop MSCC Start / MSCC Stop"],
        ["Digi audio", "mscc-virtual-audio creates VirtualA / VirtualB"],
        [
            "Config seed",
            "$HOME/.local/mscc/ only if empty (never overwrites your files later)",
        ],
        ["CAT helper", "tty0tty module → often /dev/tnt0 (after build / reboot)"],
    ],
    [1.8, 5.2],
)

doc.add_heading("3) Install the GUI setup wizard (recommended)", level=3)
add_code("sudo apt install -y ./mscc-init-gui_1.0.10_all.deb")
doc.add_paragraph("(Use the real filename if the version differs.)")
doc.add_paragraph(
    "Adds menu MSCC Init. Needs desktop packages such as python3-tk and python3-pyaudio "
    "(usually installed automatically)."
)

doc.add_heading("4) Log out / in once (first install)", level=3)
doc.add_paragraph(
    "If the install added groups (dialout, plugdev, audio) or audio services: "
    "log out and log back in, or reboot once."
)
doc.add_paragraph("Then check digi sinks:")
add_code("pactl list short sinks | grep Virtual")
doc.add_paragraph("If empty:")
add_code(
    "systemctl --user enable --now mscc-virtual-audio\n"
    "# or run once:  mscc-virtual-audio"
)

add_table(
    ["Role", "Device name"],
    [
        ["Digital speaker (recv digi out)", "VirtualA"],
        ["Digital mic (trans digi in)", "VirtualB.monitor"],
        ["WSJT Input (hear radio)", "VirtualA.monitor"],
        ["WSJT Output (TX audio)", "VirtualB"],
    ],
    [3.2, 3.8],
)

doc.add_heading("5) Configure this Pi (required once)", level=3)
p = doc.add_paragraph()
bold_run(p, "Preferred (point and click):")
for t in [
    "Applications menu → Sound & Video (typical).",
    "Click MSCC Init.",
]:
    doc.add_paragraph(t, style="List Number")

doc.add_paragraph("The wizard:")
for t in [
    "Offers to stop servers if they are already running (safe if you refuse and exit).",
    "Keyer, CAT port / PTT pin.",
    "Operator speaker and microphone — pick your phones / HAT. Any sample rate is OK "
    "(48 kHz, 96 kHz, etc.). Do not pick Proficio / Multus I/Q for headphones or mic. "
    "On Pi 4, prefer an I²S HAT or Pi headphones over a USB headset (see Operator audio hardware).",
    "Digi is fixed: VirtualA / VirtualB.monitor (not a menu choice).",
    "After save: Yes/No to start the MSCC servers.",
    "Sets operator hardware mixer levels toward full open (100%) so the MSCC client Volume can control loudness.",
]:
    doc.add_paragraph(t, style="List Bullet")

p = doc.add_paragraph()
bold_run(p, "CLI alternative")
p.add_run(" (SSH or no desktop):")
add_code("mscc-init")

doc.add_paragraph("Config files live under:")
add_code("$HOME/.local/mscc/")
doc.add_paragraph("Upgrades do not overwrite existing config.")

doc.add_heading("6) Start and stop the servers (everyday)", level=3)
p = doc.add_paragraph()
bold_run(p, "Preferred (point and click):")

add_table(
    ["Menu (usually Sound & Video)", "What it does"],
    [
        [
            "MSCC Start",
            "Starts digi sinks (best effort) + starts the three servers. "
            "Also re-applies operator speaker/mic ALSA levels to full "
            "(so phones stay loud after reboot).",
        ],
        ["MSCC Stop", "Stops the three servers."],
        ["MSCC Init", "Configure again (stops servers first if needed)."],
    ],
    [2.4, 4.6],
)

doc.add_paragraph(
    "You may get a desktop notification if libnotify-bin is installed."
)
p = doc.add_paragraph()
bold_run(p, "CLI alternative:")
add_code("mscc start\nmscc status\nmscc stop\nmscc restart")
doc.add_paragraph(
    "Then use the MSCC client on your PC to operate the radio "
    "(Pi IP / hostname as configured)."
)

# --- Phones vs digi ---
doc.add_heading("Operator phones vs digi (what to expect)", level=2)

doc.add_heading("Phones (operator path)", level=3)
for t in [
    "Proficio radio I/Q is always 96 kHz on the wire.",
    "Your operator device can be 96 kHz or not (e.g. 48 kHz). The servers handle the difference.",
    "On Pi 4, best results with an I²S HAT or Pi jack — not a second USB sound card "
    "(see Operator audio hardware).",
    "Volume after reboot: run MSCC Start (or mscc start). Levels are re-applied automatically.",
    "If still quiet: re-run MSCC Init, or raise levels once with the desktop volume control / "
    "alsamixer for that sound card.",
    "Day-to-day loudness: use the MSCC client Volume control.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("Digi (WSJT etc.)", level=3)
doc.add_paragraph("Typical WSJT-X on the Pi:")
add_table(
    ["WSJT setting", "Value"],
    [
        ["Input", "VirtualA.monitor"],
        ["Output", "VirtualB"],
        ["CAT", "As set in MSCC Init (often /dev/tnt0)"],
    ],
    [2.2, 4.8],
)

doc.add_paragraph(
    "Digi uses VirtualA / VirtualB (software). That does not add a second USB sound device, "
    "so it is a good fit for Pi 4 as well as Pi 5."
)
doc.add_paragraph(
    "Digi levels use Pulse/PipeWire, not the same as phone alsamixer."
)
doc.add_paragraph(
    "If WSJT’s receive bar is weak while the MSCC client is already loud:"
)
add_code("pactl set-sink-volume VirtualA 100%\npactl set-sink-mute VirtualA 0")
doc.add_paragraph(
    "(That raises the digi path into WSJT. Then trim with MSCC client Volume if needed.)"
)
p = doc.add_paragraph()
bold_run(p, "WSJT-X and CPU (especially Pi 4): ")
p.add_run(
    "MSCC already uses real CPU for the radio. FT8 with many decoder threads can load the Pi heavily. "
    "If the desktop feels sluggish or decodes fall behind, lower WSJT’s thread count (or use single-thread "
    "decode). A Pi 5 has more headroom for MSCC + multi-thread FT8 together."
)
doc.add_paragraph(
    "Virtual sinks should return at desktop login via the mscc-virtual-audio user service. "
    "If digi vanishes after reboot:"
)
add_code("systemctl --user enable --now mscc-virtual-audio")

# --- Did it work ---
doc.add_heading("Did it work?", level=2)
add_code(
    "mscc status\n"
    "ls /dev/tnt0 /dev/tnt1\n"
    "pactl list short sinks | grep Virtual\n"
    "ldd $HOME/mscc/sdrcore-recv | grep portaudio"
)

add_table(
    ["Check", "Good sign"],
    [
        ["mscc status", "sdrcore-recv, sdrcore-trans, and ms-sdr show running"],
        ["/dev/tnt0", "Exists (CAT / PTT when tty0tty built)"],
        ["pactl … Virtual", "VirtualA / VirtualB present"],
        ["ldd … portaudio", "/usr/local/lib"],
        ["PC client", "Connects to this Pi"],
    ],
    [2.4, 4.6],
)

doc.add_paragraph("Logs (if something fails):")
for t in [
    "Often: ~/sdrcore-recv.log, ~/sdrcore-trans.log, ~/ms-sdr.log",
    "Or under: $HOME/.local/mscc/",
    "Start script text: $HOME/mscc/logs/",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph("Useful digi / audio lines:")
add_code(
    'grep -E "VirtualA|DUAL STREAM|ALSA card|portaudio|FAILED" ~/sdrcore-recv.log | tail -30'
)

# --- Everyday ---
doc.add_heading("Everyday use (after first install)", level=2)
for t in [
    "Power on Pi and radio.",
    "Log into the Pi desktop (so Pulse/PipeWire and Virtual* can start).",
    "Menu → Sound & Video → MSCC Start (or Terminal: mscc start).",
    "Use MSCC on the PC as usual.",
    "When done: MSCC Stop (optional if you leave the Pi running).",
]:
    doc.add_paragraph(t, style="List Number")

doc.add_paragraph(
    "Re-run MSCC Init only if you change phones, HAT, or CAT wiring."
)

# --- Failures ---
doc.add_heading("If something fails", level=2)
add_table(
    ["Symptom", "What to try"],
    [
        [
            "mscc: command not found",
            "New Desktop Terminal, or log out/in",
        ],
        [
            "No menu entries",
            "Confirm packages installed; look under Sound & Video; log out/in",
        ],
        [
            "mscc-init-gui missing",
            "Install mscc-init-gui_*.deb; need desktop + python3-tk / python3-pyaudio",
        ],
        [
            "No /dev/tnt0",
            "Reboot once; install headers: sudo apt install -y linux-headers-$(uname -r)",
        ],
        [
            "No VirtualA/B",
            "mscc-virtual-audio or systemctl --user enable --now mscc-virtual-audio",
        ],
        [
            "Digi devices not listed in init / WSJT",
            "Install mscc-portaudio first; check ldd → /usr/local",
        ],
        [
            "Wrong PortAudio (ldd shows only /lib/...)",
            "Install mscc-portaudio; use mscc 1.0.23+",
        ],
        [
            "Quiet phones after reboot",
            "MSCC Start (re-applies levels); re-run MSCC Init if needed",
        ],
        [
            "Quiet digi / WSJT RX weak",
            "pactl set-sink-volume VirtualA 100%; WSJT Input = VirtualA.monitor",
        ],
        [
            "Audio wrong device",
            "Re-run MSCC Init. Digi stays VirtualA / VirtualB.monitor",
        ],
        [
            "Spectrum / phones freeze; mscc status still running",
            "Often USB audio dropped under load. On Pi 4, avoid USB headset + Proficio. "
            "Use an I²S HAT or digi path. Then mscc stop / mscc start and reconnect the client. "
            "Check: dmesg -T | grep -iE 'disconnect|usb'",
        ],
        [
            "WSJT makes the Pi very busy / late decodes",
            "Lower WSJT decoder threads on Pi 4; prefer digi VirtualA/B; "
            "consider Pi 5 for heavy multi-thread FT8 + MSCC",
        ],
        [
            "Permission denied on serial/tnt",
            "Log out/in or reboot (groups dialout / plugdev)",
        ],
        [
            "Servers won’t start from menu",
            "Terminal: mscc start and read the error text",
        ],
        [
            "Client won’t connect",
            "Same network; firewall; correct Pi IP; mscc status",
        ],
        [
            "AudioInjector HAT silent after first install",
            "Reboot once if the package added the audio overlay",
        ],
    ],
    [2.8, 4.2],
)

p = doc.add_paragraph()
bold_run(
    p,
    "This stack is only for Raspberry Pi OS 64-bit on Pi 4/5. "
    "Other Linux systems are not supported.",
)

# --- Servers optional ---
doc.add_heading("What the three servers are (optional reading)", level=2)
add_table(
    ["Program", "Job"],
    [
        ["sdrcore-recv", "Receive: radio I/Q in, audio to phones and digi out"],
        ["sdrcore-trans", "Transmit: mic / digi in, radio I/Q out"],
        ["ms-sdr", "Control hub: talks to Proficio and to the PC client"],
    ],
    [2.0, 5.0],
)
doc.add_paragraph(
    "You do not start them individually for normal use — use MSCC Start / mscc start."
)

# --- Cheat sheet ---
doc.add_heading("One-line cheat sheet", level=2)
add_code(
    "Desktop Terminal (order matters):\n"
    "  1. sudo apt install -y ./mscc-portaudio_*_arm64.deb\n"
    "  2. sudo apt install -y ./mscc_*_arm64.deb\n"
    "  3. sudo apt install -y ./mscc-init-gui_*_all.deb\n"
    "  (log out/in if first time)\n"
    "  Optional checks:\n"
    "     ldconfig -p | grep portaudio\n"
    "     ldd $HOME/mscc/sdrcore-recv | grep portaudio\n"
    "     pactl list short sinks | grep Virtual\n"
    "\n"
    "Menu → Sound & Video:\n"
    "  MSCC Init  →  configure once  →  Yes to start (or MSCC Start later)\n"
    "  MSCC Start / MSCC Stop day to day\n"
    "\n"
    "PC: run MSCC client → connect to Pi\n"
    "\n"
    "WSJT (on Pi): Input = VirtualA.monitor   Output = VirtualB\n"
    "\n"
    "Pi 4 operator audio: I²S HAT or Pi jack preferred (not USB headset + Proficio)\n"
    "Pi 4 + FT8: keep WSJT thread count modest"
)

doc.add_paragraph("That’s it.")

doc.add_heading("For developers / handoff", level=2)
doc.add_paragraph(
    "See README.md (package build) and resume.md (status, audio/pan contracts, open items). "
    "Do not put long engineering notes in this operator card — keep this file short and one path."
)

doc.save(out)
print("Wrote", out)

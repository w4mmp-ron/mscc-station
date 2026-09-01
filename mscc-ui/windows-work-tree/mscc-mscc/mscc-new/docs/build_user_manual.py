"""Generate MSCC WPF User Manual (.docx) for Multus SDR group release."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).resolve().parent / "MSCC-WPF-User-Manual.docx"


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_caption_placeholder(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[SCREENSHOT: {caption}]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = note.add_run("(Insert your capture here — crop to the relevant panel if possible.)")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_table(doc: Document, headers, rows) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A3A6B")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
        if ri % 2 == 1:
            for ci in range(len(headers)):
                set_cell_shading(t.rows[ri + 1].cells[ci], "F0F4FA")
    doc.add_paragraph()


def p(doc, text="", bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for r in para.runs:
        r.font.size = Pt(11)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(text, style="List Number")
    for r in para.runs:
        r.font.size = Pt(11)
    return para


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    for i in range(1, 4):
        h = styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
        h.font.size = Pt(18 if i == 1 else 14 if i == 2 else 12)

    # Header / footer
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.text = "Multus SDR  ·  MSCC WPF User Manual"
    for r in hp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("MSCC WPF  ·  Client 8.4.x  ·  Page ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run2 = fp.add_run()
    run2._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run3 = fp.add_run()
    run3._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run4 = fp.add_run()
    run4._r.append(fld_end)
    for r in fp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MSCC")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Multus Software Control Console")
    r.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("WPF Client User Manual")
    r.bold = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "For Multus SDR Proficio / Geminus operators\n"
        "Software version: MSCC WPF 8.4.x (and later 8.x)\n"
        "Prepared for Multus SDR group release"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    p(
        doc,
        "Screenshot placeholders appear as [SCREENSHOT: …]. Replace them with your own captures; "
        "crop to the panel being described when possible.",
        italic=True,
        size=10,
    )

    doc.add_page_break()

    # Contents
    doc.add_heading("Contents", level=1)
    for item in [
        "1. Introduction",
        "2. System requirements",
        "3. Installation and first-time setup",
        "4. Starting and stopping a session",
        "5. Main window overview",
        "6. Button appearance (selected vs not selected)",
        "7. Left panel — connection, audio, filters",
        "8. Top panel — meters, VFOs, band bar",
        "9. Right panel — operate controls",
        "10. Tabs reference",
        "11. Frequency calibration (FREQ CAL)",
        "12. Noise tools: NB, NR, AN",
        "13. Digital / DIG-U notes",
        "14. Configuration files and reset",
        "15. Troubleshooting",
        "16. Glossary",
    ]:
        bullet(doc, item)

    doc.add_page_break()

    # 1
    doc.add_heading("1. Introduction", level=1)
    p(
        doc,
        "MSCC (Multus Software Control Console) is the PC application used to operate Multus SDR radios "
        "such as the Proficio and Geminus. This WPF edition is the modern Windows client: it talks to the "
        "radio stack over UDP and shows spectrum, meters, and day-to-day controls.",
    )

    doc.add_heading("What the client talks to", level=2)
    p(doc, "A normal session involves three backend programs plus the radio firmware:")
    bullet(doc, "ms-sdr (or ms-sdr-MKII) — command hub, USB to the radio, keep-alive")
    bullet(doc, "mscc-recv (sdrcore-recv) — receive DSP, spectrum/waterfall, NB/NR/AN")
    bullet(doc, "Mscc-trans (sdrcore-trans) — transmit audio / digi path")
    bullet(doc, "MSCC.Wpf.exe — this GUI")
    p(
        doc,
        "On a single PC you usually let MSCC launch the backends (Launch Servers checked). "
        "Advanced users can run backends separately and connect the GUI only.",
    )
    add_caption_placeholder(doc, "Main window — full overview")

    # 2
    doc.add_heading("2. System requirements", level=1)
    add_table(
        doc,
        ["Item", "Requirement"],
        [
            ["OS", "Windows 10 or 11 (64-bit recommended)"],
            ["Runtime", ".NET 9 Desktop Runtime (Windows), if not bundled with your installer"],
            ["USB", "Radio connected and recognized (Proficio / Geminus USB)"],
            ["Audio", "Operator speaker + microphone configured (SETTINGS tab)"],
            ["COM port", "Configured when required by your station / CAT path"],
            ["Install folder", "Typical: C:\\mscc-net9\\ with client + server binaries together"],
            ["Config data", "%LocalAppData%\\MSCC-NET9\\ (created automatically on first run)"],
        ],
    )

    # 3
    doc.add_heading("3. Installation and first-time setup", level=1)
    doc.add_heading("3.1 Install layout", level=2)
    p(doc, "Your Multus package should place at least these items in one folder (example C:\\mscc-net9\\):")
    bullet(doc, "MSCC.Wpf.exe and supporting DLLs (MSCC.Core, NAudio, and related)")
    bullet(doc, "ms-sdr-MKII.exe (or the ms-sdr binary as shipped)")
    bullet(doc, "mscc-recv.exe and Mscc-trans.exe (names as shipped)")
    bullet(doc, "init-files\\ (optional seed configs for first run / reset)")
    p(
        doc,
        "Keeping the GUI and subsystem binaries in the same folder lets the client discover backends "
        "the same way Multus packaging expects.",
    )

    doc.add_heading("3.2 First launch checklist", level=2)
    numbered(doc, "Connect the radio USB and power it on.")
    numbered(doc, "Start MSCC.Wpf.exe.")
    numbered(doc, "Open the SETTINGS tab.")
    numbered(doc, "Set the COM port used by your station (if required).")
    numbered(doc, "Set Operator speaker and Operator microphone.")
    numbered(
        doc,
        "Confirm Server address is 127.0.0.1 for a local radio PC (or your remote host if backends run elsewhere).",
    )
    numbered(doc, "Leave Launch Servers checked for a normal single-PC station.")
    numbered(doc, "Press Start (connection area).")
    add_caption_placeholder(doc, "SETTINGS tab — COM port and audio devices")

    doc.add_heading("3.3 Launch Servers vs connect-only", level=2)
    add_table(
        doc,
        ["Option", "When to use"],
        [
            ["Launch Servers ON", "Normal home station: MSCC starts ms-sdr, recv, and trans for you."],
            [
                "Launch Servers OFF",
                "Backends already running (batch file, second PC, or headless appliance). GUI only connects.",
            ],
            ["Auto Start", "Starts the session when the MSCC window loads (saved preference)."],
        ],
    )
    p(
        doc,
        "When Launch Servers is ON and you press Stop (or exit after launching backends), MSCC sends a stop "
        "so those processes shut down cleanly. When Launch Servers was OFF, Stop/exit leave backends running.",
    )

    # 4
    doc.add_heading("4. Starting and stopping a session", level=1)
    doc.add_heading("4.1 Start", level=2)
    bullet(doc, "Press Start when COM and operator audio requirements are satisfied.")
    bullet(
        doc,
        "MSCC waits briefly for backends, claims a session with ms-sdr, then shows firmware (FW) and core versions.",
    )
    bullet(doc, "Only one GUI client may own a session at a time.")

    doc.add_heading("4.2 Stop", level=2)
    bullet(doc, "Press Stop to end the radio session.")
    bullet(doc, "If this app launched the backends, they are stopped as well.")

    doc.add_heading("4.3 Keep-alive", level=2)
    p(
        doc,
        "While connected, MSCC and ms-sdr exchange keep-alive messages. If the GUI does not hear the server "
        "for about 10 seconds after a short startup grace period, a warning appears. Choose Continue only if "
        "you know the stack is recovering; otherwise Stop and Start again.",
    )

    # 5
    doc.add_heading("5. Main window overview", level=1)
    p(doc, "Think of the window in four regions:")
    add_table(
        doc,
        ["Region", "What you find there"],
        [
            [
                "Left",
                "Start/Stop, server IP/port, Auto Start, Launch Servers, Phones/Digital audio, "
                "Lo/Hi/CW filters, zoom, RIT tools, temperatures",
            ],
            [
                "Top / center-top",
                "S-meter, ALC, VFO A/B, frequency, radio model & band buttons, GEN",
            ],
            [
                "Center tabs",
                "MAIN spectrum, CW, RX/TX, Favorites, QRP CAL, AMP CAL, RX IQ, TX IQ, FREQ CAL, SETTINGS",
            ],
            [
                "Right",
                "Always-visible operate cluster: PTT, TUN, AMP, AGC, COMP, power, modes, "
                "NB/NR/AN, MON, TX BW, CW speed/pitch, LOG, versions",
            ],
        ],
    )
    add_caption_placeholder(doc, "Annotated main window regions (left / top / tabs / right)")

    # 6
    doc.add_heading("6. Button appearance (selected vs not selected)", level=1)
    p(doc, "Most gold toggle buttons share one visual language across MSCC:")
    add_table(
        doc,
        ["State", "Appearance", "Meaning"],
        [
            ["Not selected (off)", "Bright gold face, black text", "Feature or mode is off / not chosen"],
            ["Selected (on)", "Darker gold face, white text", "Feature or mode is on / active"],
            [
                "Alert latched (e.g. PTT / TUN / AMP)",
                "Red-style alert face",
                "Transmit or amp path engaged — treat with care",
            ],
            ["Disabled", "Grayed / reduced opacity", "Not available for current model or mode"],
        ],
    )
    p(
        doc,
        "Examples: mode buttons (USB, LSB, CW…), NB / NR / AN, AMP, COMP, AGC cycle, and similar toggles.",
    )

    # 7
    doc.add_heading("7. Left panel — connection, audio, filters", level=1)
    doc.add_heading("7.1 Connection", level=2)
    bullet(doc, "Server address / port — usually 127.0.0.1 for a local station PC.")
    bullet(doc, "Start / Stop — session control.")
    bullet(doc, "Auto Start — start when the window opens.")
    bullet(doc, "Launch Servers — spawn backends with Start.")

    doc.add_heading("7.2 Audio path P / D", level=2)
    bullet(doc, "P (Phones) — operator speaker and mic.")
    bullet(doc, "D (Digital) — digital / VAC path for apps such as WSJT-X.")
    p(doc, "DIG-U mode pairs naturally with Digital audio.")

    doc.add_heading("7.3 Filters and zoom", level=2)
    bullet(doc, "Low cut / High cut — cycle filter edges for voice-style modes.")
    bullet(doc, "CW bandwidth — cycle CW filter widths.")
    bullet(doc, "Step — VFO tune step.")
    bullet(
        doc,
        "Zoom — display zoom of the panadapter (1×–4×) around the VFO; magnifies the view of the same RF data.",
    )

    # 8
    doc.add_heading("8. Top panel — meters, VFOs, band bar", level=1)
    doc.add_heading("8.1 Meters", level=2)
    bullet(doc, "S-meter — receive signal strength.")
    bullet(doc, "ALC — transmit ALC indication.")

    doc.add_heading("8.2 VFO A / VFO B", level=2)
    bullet(doc, "Click a VFO panel to make it active (mouse wheel tunes the active VFO).")
    bullet(doc, "RIT offset may be shown per VFO when used.")

    doc.add_heading("8.3 Radio model button", level=2)
    p(
        doc,
        "Cycles Proficio ↔ Geminus (UI band enablement and waterfall banks). Match the model to your hardware:",
    )
    bullet(doc, "Proficio — HF bands; GEN rotates standard HF time/frequency beacons (WWV / CHU / RWM / USER style).")
    bullet(
        doc,
        "Geminus — enables 2200 m / 630 m; GEN uses LF carriers useful for frequency calibration.",
    )
    add_caption_placeholder(doc, "Band bar with Proficio/Geminus model button and GEN")

    doc.add_heading("8.4 Band buttons", level=2)
    p(
        doc,
        "Click a band to QSY to a default or last-used frequency for that band. Gray (disabled) bands "
        "cannot be used for the selected radio model.",
    )

    # 9
    doc.add_heading("9. Right panel — operate controls", level=1)
    p(doc, "These stay visible while you use tabs, so you can work spectrum and still PTT or change mode.")

    doc.add_heading("9.1 Transmit cluster", level=2)
    add_table(
        doc,
        ["Control", "Function"],
        [
            ["PTT", "Push-to-talk (alert style when on)."],
            ["TUN", "Tune / carrier helper for calibration and matching."],
            ["AMP", "External amplifier path / PA control (alert when engaged)."],
            [
                "Drive / power",
                "Context-sensitive drive: TUN, CW, SSB, or AM carrier depending on mode.",
            ],
        ],
    )

    doc.add_heading("9.2 AGC, compression, spectrum popup", level=2)
    add_table(
        doc,
        ["Control", "Function"],
        [
            ["AGC (SLO / MED / FAST)", "Cycles AGC speed."],
            ["AGC fast release", "Fine timing for AGC release (ms)."],
            ["COMP", "Compression on/off (phones path)."],
            ["Compression level", "dB amount when COMP is active (phones / P)."],
            ["S/W", "Opens Spectrum / Waterfall settings popup."],
        ],
    )

    doc.add_heading("9.3 Modes", level=2)
    bullet(doc, "LSB, USB, DIG-U, CW, AM — FM may appear disabled until implemented.")
    bullet(doc, "DIG-U uses its own low/high cut memory and expects Digital (D) audio for many digi apps.")

    doc.add_heading("9.4 Noise tools and CW face controls", level=2)
    add_table(
        doc,
        ["Control", "Function"],
        [
            ["NB", "Noise Blanker on/off. Pulse and threshold sliders below when used."],
            ["NR", "Noise Reduction on/off. Level slider when on."],
            ["AN", "Auto Notch on/off (adaptive tone notch). Selected = on = darker gold."],
            ["MON", "Monitor."],
            ["TX BW", "Transmit bandwidth cycle."],
            ["CW speed / pitch", "Quick CW face controls (full set also on CW tab)."],
            ["LOG", "Debug / monitor log window."],
            ["Versions", "MSCC (this client), Core (ms-sdr), FW (radio firmware)."],
        ],
    )
    add_caption_placeholder(doc, "Right operate panel — modes, NB, NR, AN")

    # 10
    doc.add_heading("10. Tabs reference", level=1)

    doc.add_heading("10.1 MAIN", level=2)
    p(
        doc,
        "Primary spectrum and waterfall display. Click in the panadapter to tune (subject to Auto Snap "
        "settings from the S/W popup). Most operating is done with MAIN visible plus the right panel.",
    )
    add_caption_placeholder(doc, "MAIN tab with spectrum and waterfall")

    doc.add_heading("10.2 CW", level=2)
    p(
        doc,
        "Electronic keyer and CW parameters: speed, pitch, keyer mode, spacing, paddle, weight, hold, "
        "QSK, phones, and related options. Use when setting up CW beyond the quick right-panel controls.",
    )
    add_caption_placeholder(doc, "CW tab")

    doc.add_heading("10.3 RX/TX", level=2)
    p(
        doc,
        "Receive/transmit preferences and power levels (tune / CW / SSB / AM). NB, NR, and AN also appear "
        "on the right operate panel for quick access while watching the spectrum.",
    )
    add_caption_placeholder(doc, "RX/TX tab — power levels")

    doc.add_heading("10.4 FAVORITES", level=2)
    p(doc, "Client-only named memories per band. Example: save “FT8” on 20 m and a separate “FT8” on 15 m.")
    bullet(doc, "Choose the band context (or use the radio’s current band).")
    bullet(doc, "Enter a name → Save current conditions.")
    bullet(doc, "Select a row → Recall or Delete.")
    p(doc, "Favorites are stored under %LocalAppData%\\MSCC-NET9\\ (MSCC_Favorites.ini).")

    doc.add_heading("10.5 QRP CAL", level=2)
    p(
        doc,
        "Transceiver / QRP power calibration. Requires AMP off. Select band and use the calibrate workflow "
        "and drive controls as labeled. Band lamps show client-side calibrated status for convenience. "
        "Follow Multus hardware notes for dummy load and power meter use.",
    )

    doc.add_heading("10.6 AMP CAL", level=2)
    p(
        doc,
        "External amplifier calibration. Requires AMP on. Band-oriented workflow with amp-specific steps. "
        "Grayed when AMP is off.",
    )

    doc.add_heading("10.7 RX IQ", level=2)
    p(
        doc,
        "Receive I/Q balance tools: start session, optional +24 kHz LO aid, offset, commit, apply-all, "
        "LO fine offset. Use carefully; incorrect I/Q settings can hurt image rejection.",
    )

    doc.add_heading("10.8 TX IQ", level=2)
    p(
        doc,
        "Transmit I/Q balance (QRP / AMP off). Manual offset with tune carrier; commit when ready. "
        "Factory reset of I/Q tables is available with confirmation.",
    )

    doc.add_heading("10.9 FREQ CAL", level=2)
    p(doc, "See the dedicated section below.")

    doc.add_heading("10.10 SETTINGS", level=2)
    bullet(doc, "UI appearance (colors / chrome).")
    bullet(doc, "COM port.")
    bullet(doc, "Audio devices (operator and digital).")
    bullet(doc, "Optional external Wi-Fi SWR meter.")
    bullet(
        doc,
        "Configuration reset — reseeds configs from install init-files when present (does not wipe logs by design).",
    )

    # 11
    doc.add_heading("11. Frequency calibration (FREQ CAL)", level=1)
    p(doc, "Use FREQ CAL with a known standard (WWV, CHU, or LF carriers on the Geminus GEN list). Typical controls:")
    add_table(
        doc,
        ["Button", "Purpose"],
        [
            ["LOOSE", "Widens the peak-search window for a coarse check."],
            ["AUTO", "Runs automatic calibrate sequence (coarse then fine stages)."],
            ["MANUAL", "Enters manual PPM adjustment mode."],
            ["CHECK", "Measures error near the current frequency without a full auto run."],
            ["RESET", "Resets calibration state as implemented."],
            ["PPM − / +", "Manual PPM steps (rate-limited to protect the radio)."],
        ],
    )
    p(doc, "Tips:")
    bullet(doc, "Use a strong, known carrier; weak peaks can fail CHECK.")
    bullet(doc, "Prefer GEN standards when available.")
    bullet(doc, "Allow AUTO to finish; watch the status label and progress bar.")
    add_caption_placeholder(doc, "FREQ CAL tab")

    # 12
    doc.add_heading("12. Noise tools: NB, NR, AN", level=1)
    add_table(
        doc,
        ["Tool", "Name", "Selected (darker) means", "Notes"],
        [
            ["NB", "Noise Blanker", "Blanker ON", "Pulse width and threshold sliders below the buttons."],
            ["NR", "Noise Reduction", "Reduction ON", "Level slider when NR is on."],
            [
                "AN",
                "Auto Notch",
                "Notch ON",
                "Adaptive notch for steady tones/heterodynes; no extra level control.",
            ],
        ],
    )
    p(
        doc,
        "Visual rule (same as other toggles): bright gold = off; darker gold with white text = on. "
        "When AN is selected (on), the client sends auto-notch enable = 1 to the receive DSP.",
    )

    # 13
    doc.add_heading("13. Digital / DIG-U notes", level=1)
    numbered(doc, "Select DIG-U mode.")
    numbered(doc, "Set audio path to D (Digital).")
    numbered(
        doc,
        "Point your digi app (for example WSJT-X) at the Multus digital devices configured in SETTINGS.",
    )
    numbered(doc, "Confirm TX drive on the right-panel power control for digi.")
    p(
        doc,
        "Exact device names depend on your PortAudio / virtual audio setup and Multus install package.",
    )

    # 14
    doc.add_heading("14. Configuration files and reset", level=1)
    p(doc, "Client preferences live primarily in:")
    p(doc, "%LocalAppData%\\MSCC-NET9\\", bold=True)
    add_table(
        doc,
        ["File (examples)", "Role"],
        [
            ["MSCC_Client.ini", "UI, connection, power, spectrum/waterfall prefs"],
            ["MSCC_LastUsed.ini / VFOB", "Per-band last frequencies"],
            ["MSCC_Favorites.ini", "Favorites"],
            ["user_controls.ini", "Shared control snapshot used with ms-sdr"],
            ["startup.ini", "Last frequency/mode for start / appliance style resume"],
            ["cw.ini", "CW keyer parameters for ms-sdr"],
            ["freq_cal.ini / power_cal.ini / …", "Calibration-related data"],
            ["logs\\", "mscc.log, ms-sdr.log, sdrcore-recv/trans logs"],
        ],
    )
    p(
        doc,
        "SETTINGS → configuration reset can delete client configs and reseed from install init-files, "
        "then restart MSCC. Back up the folder first if you care about custom calibrations or favorites.",
    )

    # 15
    doc.add_heading("15. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "What to try"],
        [
            ["Start disabled / setup incomplete", "Set COM port and operator speaker/mic in SETTINGS."],
            [
                "No spectrum / no sound",
                "Confirm Launch Servers, radio USB, and backends in Task Manager; check logs\\.",
            ],
            [
                "Keep-alive lost",
                "Stop, wait a few seconds, Start again; ensure only one client; check USB cable/power.",
            ],
            [
                "Second PC cannot connect",
                "Only one session owner; stop the first client or run intentional connect-only design.",
            ],
            ["Wrong bands enabled", "Set radio model (Proficio vs Geminus) on the band bar."],
            [
                "FREQ CAL CHECK fails",
                "Stronger standard signal; try LOOSE; center near the carrier; try AUTO.",
            ],
            [
                "AN seems wrong",
                "Selected (darker) = ON. Confirm sdrcore-recv log shows Enabled: 1 when selected.",
            ],
            ["After COM/audio change", "Stop, then Start so backends pick up new devices."],
        ],
    )
    p(
        doc,
        "Log locations: %LocalAppData%\\MSCC-NET9\\logs\\ — mscc.log (GUI), ms-sdr.log, "
        "sdrcore-recv.log, sdrcore-trans.log.",
    )

    # 16
    doc.add_heading("16. Glossary", level=1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["MSCC", "Multus Software Control Console (this GUI)."],
            ["ms-sdr", "Main server process bridging GUI, DSP cores, and radio USB."],
            ["recv / trans", "Receive and transmit DSP cores."],
            ["FW", "Radio firmware version reported at connect."],
            ["Core", "ms-sdr version string."],
            ["QRP CAL", "Transceiver power calibration tab."],
            ["GEN", "General coverage / standard frequency list button."],
            ["VAC", "Virtual audio cable (or equivalent) for digi apps."],
            ["AN", "Auto Notch."],
        ],
    )

    doc.add_heading("Support and community", level=1)
    p(
        doc,
        "For Multus SDR group releases: include the version shown on the right panel as MSCC: x.y.z when "
        "asking questions. Attach relevant log snippets (not necessarily full multi-megabyte files) when "
        "reporting connect or keep-alive issues.",
    )
    p(
        doc,
        "Document prepared for group distribution. Hardware calibration procedures remain subject to Multus "
        "hardware documentation and safe station practice (dummy loads, power limits, amplifier interlocks).",
        italic=True,
        size=10,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Size {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
